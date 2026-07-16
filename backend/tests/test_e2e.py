from io import BytesIO
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from app.main import app
from app.database import SessionLocal
from app.models import Task
from app.services import run_task

def test_docx_to_cited_artifact():
    client=TestClient(app)
    workspace=client.post("/api/workspaces",json={"name":"端到端测试"}).json()
    stream=BytesIO(); doc=WordDocument(); doc.add_heading("采购合同",0); doc.add_paragraph("合同金额为人民币 120000 元，截止日期为 2026 年 12 月 31 日。"); doc.save(stream)
    upload=client.post(f"/api/workspaces/{workspace['id']}/documents",files=[("files",("contract.docx",stream.getvalue(),"application/vnd.openxmlformats-officedocument.wordprocessingml.document"))])
    assert upload.status_code==201 and upload.json()[0]["status"]=="ready"
    created=client.post("/api/tasks",json={"workspace_id":workspace["id"],"message":"合同金额"})
    assert created.status_code==202
    with SessionLocal() as db:
        task=db.get(Task,created.json()["id"]); task.status="running"; db.commit(); run_task(db,task); db.refresh(task)
        assert task.status=="completed"
        assert task.final_answer_json["citations"]
        assert task.final_answer_json["artifact_ids"]
    assert client.delete(f"/api/workspaces/{workspace['id']}").status_code==204
