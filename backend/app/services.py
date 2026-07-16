import hashlib, html, json
from pathlib import Path
from sqlalchemy import select, func
from docx import Document as WordDocument
from openpyxl import Workbook, load_workbook
from .config import settings
from .models import Task, TaskEvent, Document, DocumentChunk, Artifact, now
from .schemas import TaskResult, Finding, Citation
from .security import safe_child, sha256_file
from .retrieval import hybrid_chunk_ids

def add_event(db, task_id: str, kind: str, payload: dict):
    seq = db.scalar(select(func.coalesce(func.max(TaskEvent.sequence), 0)).where(TaskEvent.task_id == task_id)) + 1
    event = TaskEvent(task_id=task_id, sequence=seq, type=kind, payload_json=payload); db.add(event); db.commit(); return event

def search_chunks(db, workspace_id: str, query: str, limit: int = 10):
    docs = {d.id:d for d in db.scalars(select(Document).where(Document.workspace_id == workspace_id, Document.status == "ready"))}
    if not docs: return []
    ids=hybrid_chunk_ids(workspace_id,query,limit)
    rows={c.id:c for c in db.scalars(select(DocumentChunk).where(DocumentChunk.id.in_(ids)))} if ids else {}
    return [(len(ids)-i,rows[cid],docs[rows[cid].document_id]) for i,cid in enumerate(ids) if cid in rows and rows[cid].document_id in docs]

def run_task(db, task: Task):
    task.plan_json = ["确认可用文档", "检索相关内容", "组织带来源结论", "生成并验证报告"]
    db.commit(); add_event(db, task.id, "plan", {"steps": task.plan_json})
    add_event(db, task.id, "step", {"title":"检索文档", "status":"running"})
    if settings.dashscope_api_key:
        from .agent import invoke_agent
        result=invoke_agent(task)
        known={c.id for c in result.citations}
        if any(cid not in known for f in result.findings for cid in f.citation_ids): raise RuntimeError("INVALID_CITATION_REFERENCE")
        for c in result.citations: add_event(db,task.id,"citation",c.model_dump())
        artifact=create_report(db,task,result); result.artifact_ids.append(artifact.id)
        task.final_answer_json=result.model_dump(); task.status="completed"; task.completed_at=now(); db.commit()
        add_event(db,task.id,"artifact",{"id":artifact.id,"name":artifact.name,"type":artifact.type}); add_event(db,task.id,"step",{"title":"任务完成","status":"completed"}); return
    hits = search_chunks(db, task.workspace_id, task.user_message)
    if not hits: raise RuntimeError("NO_RELEVANT_CONTENT")
    citations=[]; findings=[]
    for _, chunk, doc in hits[:6]:
        cid=f"c{len(citations)+1}"
        quote=chunk.content[:260].strip()
        citations.append(Citation(id=cid, document_id=doc.id, document_name=doc.original_name, location=chunk.location_json, quote=quote))
        findings.append(Finding(statement=quote, citation_ids=[cid]))
        add_event(db, task.id, "citation", {"id":cid,"document_name":doc.original_name,"location":chunk.location_json,"quote":quote})
    warnings=[]
    if not settings.dashscope_api_key: warnings.append("未配置 DASHSCOPE_API_KEY，本次使用本地检索摘要；配置后可启用 qwen3.7-plus 深度推理。")
    summary=f"已从 {len({x.document_id for x in citations})} 份文档中找到 {len(citations)} 条相关证据。以下内容按来源整理，请结合原文位置复核关键决策。"
    result=TaskResult(summary=summary, findings=findings, citations=citations, warnings=warnings, artifact_ids=[])
    artifact=create_report(db, task, result)
    result.artifact_ids.append(artifact.id)
    task.final_answer_json=result.model_dump(); task.status="completed"; task.completed_at=now(); db.commit()
    add_event(db, task.id, "artifact", {"id":artifact.id,"name":artifact.name,"type":artifact.type})
    add_event(db, task.id, "step", {"title":"任务完成", "status":"completed"})

def create_report(db, task: Task, result: TaskResult):
    root=safe_child(settings.app_data_dir,"workspaces",task.workspace_id,"artifacts",task.id); root.mkdir(parents=True,exist_ok=True)
    md=["# 文档分析报告", "", result.summary, "", "## 发现"]
    for f in result.findings: md += ["", f"- {f.statement} **[{', '.join(f.citation_ids)}]**"]
    md += ["", "## 来源"]
    for c in result.citations: md += ["", f"- **{c.id} · {c.document_name}** `{json.dumps(c.location, ensure_ascii=False)}`", f"  {c.quote}"]
    if result.warnings: md += ["", "## 警告"] + [f"- {x}" for x in result.warnings]
    path=root/"report.md"; path.write_text("\n".join(md),encoding="utf-8")
    artifact=Artifact(task_id=task.id,name="文档分析报告.md",type="markdown",path=str(path.relative_to(settings.app_data_dir)),size_bytes=path.stat().st_size,sha256=sha256_file(path)); db.add(artifact); db.commit()
    if path.read_text(encoding="utf-8").strip().startswith("#") is False: raise RuntimeError("ARTIFACT_VALIDATION_FAILED")
    return artifact

def create_docx_artifact(db, task, title, sections):
    root=safe_child(settings.app_data_dir,"workspaces",task.workspace_id,"artifacts",task.id); root.mkdir(parents=True,exist_ok=True)
    path=root/"report.docx"; doc=WordDocument(); doc.add_heading(title,0)
    for section in sections: doc.add_heading(section.get("title","章节"),1); doc.add_paragraph(section.get("content",""))
    doc.save(path); WordDocument(path)
    return _artifact(db,task,path,"Word 报告.docx","docx")

def create_xlsx_artifact(db, task, sheets):
    root=safe_child(settings.app_data_dir,"workspaces",task.workspace_id,"artifacts",task.id); root.mkdir(parents=True,exist_ok=True)
    path=root/"analysis.xlsx"; wb=Workbook(); wb.remove(wb.active)
    for spec in sheets:
        ws=wb.create_sheet(str(spec.get("name","数据"))[:31]); [ws.append(row) for row in spec.get("rows",[])]
    wb.save(path); load_workbook(path,read_only=True).close()
    return _artifact(db,task,path,"分析结果.xlsx","xlsx")

def _artifact(db,task,path,name,kind):
    a=Artifact(task_id=task.id,name=name,type=kind,path=str(path.relative_to(settings.app_data_dir)),size_bytes=path.stat().st_size,sha256=sha256_file(path)); db.add(a); db.commit(); return a
