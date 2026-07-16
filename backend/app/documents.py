from pathlib import Path
import fitz
from docx import Document as WordDocument
from openpyxl import load_workbook
from .models import DocumentChunk

def _slice(text: str, size: int = 1800):
    text = text.strip()
    return [text[i:i+size] for i in range(0, len(text), size)] or []

def parse_document(doc, path: Path) -> list[DocumentChunk]:
    suffix = Path(doc.original_name).suffix.lower(); chunks = []
    if suffix == ".pdf":
        with fitz.open(path) as pdf:
            doc.page_count = pdf.page_count
            for page_no, page in enumerate(pdf, 1):
                for part, text in enumerate(_slice(page.get_text("text"))):
                    chunks.append(_chunk(doc.id, len(chunks), text, {"type":"pdf","page":page_no,"part":part+1}))
    elif suffix == ".docx":
        word = WordDocument(path); heading = []
        for number, p in enumerate(word.paragraphs, 1):
            if p.style and p.style.name.startswith("Heading") and p.text.strip(): heading = [p.text.strip()]
            for text in _slice(p.text): chunks.append(_chunk(doc.id, len(chunks), text, {"type":"docx","heading_path":heading,"paragraph":number}))
        for table_no, table in enumerate(word.tables, 1):
            text = "\n".join(" | ".join(c.text for c in row.cells) for row in table.rows)
            for part in _slice(text): chunks.append(_chunk(doc.id, len(chunks), part, {"type":"docx","table":table_no}))
        doc.page_count = None
    elif suffix == ".xlsx":
        wb = load_workbook(path, read_only=True, data_only=False)
        doc.metadata_json = {"sheets": wb.sheetnames}
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=False):
                values = ["" if c.value is None else str(c.value) for c in row]
                if any(values): rows.append((row[0].row, " | ".join(values)))
            for i in range(0, len(rows), 30):
                batch = rows[i:i+30]
                if batch:
                    chunks.append(_chunk(doc.id, len(chunks), "\n".join(x[1] for x in batch), {"type":"xlsx","sheet":ws.title,"range":f"A{batch[0][0]}:{ws.cell(batch[-1][0], ws.max_column).coordinate}"}))
        wb.close(); doc.page_count = len(wb.sheetnames)
    if not chunks: raise ValueError("DOCUMENT_HAS_NO_TEXT")
    return chunks

def _chunk(document_id, index, content, location):
    return DocumentChunk(document_id=document_id, chunk_index=index, content=content, location_json=location, token_count=max(1, len(content)//4))
